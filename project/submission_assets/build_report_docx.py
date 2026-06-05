from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_MD = ROOT / "project" / "FINAL_REPORT.md"
OUT_DOCX = ROOT / "deliverables" / "24013778김예진_기말프로젝트_보고서.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Malgun Gothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    widths = [1800, 7560] if len(table.columns) == 2 else None
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                set_cell_width(row.cells[idx], width)


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("딥러닝실습 기말 프로젝트 | 24013778 김예진")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("딥러닝 기반 반려동물 안구 질환\n의심 예측 시스템")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("딥러닝실습 기말 프로젝트 최종 보고서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.rows[0].cells[0].text = "제출자"
    meta.rows[0].cells[1].text = "24013778 김예진"
    meta.rows[1].cells[0].text = "과목"
    meta.rows[1].cells[1].text = "딥러닝실습"
    meta.rows[2].cells[0].text = "데이터"
    meta.rows[2].cells[1].text = "AIHub 반려동물 안구 질환 데이터"
    meta.rows[3].cells[0].text = "구현"
    meta.rows[3].cells[1].text = "PyTorch EfficientNet-B0, Django Web App"
    style_table(meta)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
    style_table(table)
    doc.add_paragraph()


def add_code_block(doc: Document, code: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    cell.text = "\n".join(code)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    apply_document_styles(doc)
    add_footer(doc)
    add_title_page(doc)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue
        if stripped.startswith("# "):
            idx += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("|"):
            rows, next_idx = parse_markdown_table(lines, idx)
            add_table_from_rows(doc, rows)
            idx = next_idx
            continue
        elif stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        idx += 1

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
